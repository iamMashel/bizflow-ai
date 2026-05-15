from io import BytesIO
from uuid import UUID

import pytest
from docx import Document

from app.core.config import Settings
from app.services.ingestion_service import (
    IngestedDocument,
    IngestionChunk,
    IngestionService,
    IngestionServiceError,
)

USER_ID = UUID("00000000-0000-0000-0000-000000000001")
DOCUMENT_ID = UUID("00000000-0000-0000-0000-000000000101")


class FakeEmbeddingService:
    def __init__(self) -> None:
        self.inputs: list[str] = []

    def embed_text(self, text: str) -> list[float]:
        self.inputs.append(text)
        return [0.1, 0.2, 0.3]


class FakeIngestionService(IngestionService):
    def __init__(
        self,
        *,
        filename: str = "bizflow-test.txt",
        content: bytes = b"Alpha beta.\n\nGamma delta.",
        fail_replace: bool = False,
    ) -> None:
        self.fake_embeddings = FakeEmbeddingService()
        super().__init__(
            settings=Settings(
                supabase_url="https://example.supabase.co",
                supabase_anon_key="anon",
                supabase_service_role_key="service",
            ),
            embedding_service=self.fake_embeddings,  # type: ignore[arg-type]
        )
        self.filename = filename
        self.content = content
        self.fail_replace = fail_replace
        self.statuses: list[str] = []
        self.created_chunks: list[str] = []
        self.created_metadata: list[dict[str, object]] = []

    async def _get_document(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
    ) -> IngestedDocument:
        assert document_id == DOCUMENT_ID
        assert user_id == USER_ID
        assert access_token == "test-access-token"
        return IngestedDocument(
            id=document_id,
            filename=self.filename,
            storage_path=f"{user_id}/file/{self.filename}",
        )

    async def _download_storage_object(self, storage_path: str) -> bytes:
        assert storage_path.endswith(self.filename)
        return self.content

    async def _replace_chunks(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
        chunks: list[IngestionChunk],
    ) -> None:
        if self.fail_replace:
            raise IngestionServiceError("Unable to store document chunks.")

        for chunk in chunks:
            self.fake_embeddings.embed_text(chunk.content)
        self.created_chunks = [chunk.content for chunk in chunks]
        self.created_metadata = [chunk.metadata for chunk in chunks]

    async def _update_document_status(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
        status: str,
        metadata_patch: dict[str, object] | None = None,
    ) -> None:
        _ = metadata_patch
        self.statuses.append(status)


def build_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Acme proposal overview.")
    document.add_paragraph("Implementation starts next month.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Owner"
    table.rows[0].cells[1].text = "BizFlow AI"
    document.save(buffer)
    return buffer.getvalue()


class FakePdfPage:
    def __init__(self, text: str) -> None:
        self.text = text

    def extract_text(self) -> str:
        return self.text


class FakePdfReader:
    def __init__(self, stream: BytesIO) -> None:
        _ = stream
        self.pages = [
            FakePdfPage("First page about payment terms."),
            FakePdfPage("Second page about implementation."),
        ]


@pytest.mark.asyncio
async def test_txt_ingestion_extracts_text_creates_chunks_and_completes() -> None:
    service = FakeIngestionService()

    result = await async_test_ingest(service)

    assert result.status == "completed"
    assert result.chunks_created == 1
    assert service.created_chunks == ["Alpha beta.\n\nGamma delta."]
    assert service.created_metadata == [{}]
    assert service.fake_embeddings.inputs == ["Alpha beta.\n\nGamma delta."]
    assert service.statuses == ["processing", "completed"]


@pytest.mark.asyncio
async def test_md_ingestion_is_supported() -> None:
    service = FakeIngestionService(filename="notes.md", content=b"# Heading\n\nBody")

    result = await async_test_ingest(service)

    assert result.chunks_created == 1
    assert service.statuses == ["processing", "completed"]


@pytest.mark.asyncio
async def test_docx_ingestion_extracts_text_creates_chunks_and_completes() -> None:
    service = FakeIngestionService(filename="proposal.docx", content=build_docx_bytes())

    result = await async_test_ingest(service)

    assert result.status == "completed"
    assert result.chunks_created == 1
    assert service.created_chunks == [
        "Acme proposal overview.\n\nImplementation starts next month.\n\nOwner | BizFlow AI"
    ]
    assert service.created_metadata == [{}]
    assert service.fake_embeddings.inputs == service.created_chunks
    assert service.statuses == ["processing", "completed"]


@pytest.mark.asyncio
async def test_pdf_ingestion_extracts_page_text_with_page_metadata(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.services.ingestion_service.PdfReader", FakePdfReader)
    service = FakeIngestionService(filename="proposal.pdf", content=b"%PDF test")

    result = await async_test_ingest(service)

    assert result.status == "completed"
    assert result.chunks_created == 2
    assert service.created_chunks == [
        "First page about payment terms.",
        "Second page about implementation.",
    ]
    assert service.created_metadata == [{"page_number": 1}, {"page_number": 2}]
    assert service.fake_embeddings.inputs == service.created_chunks
    assert service.statuses == ["processing", "completed"]


@pytest.mark.asyncio
async def test_failure_updates_document_status_to_failed() -> None:
    service = FakeIngestionService(fail_replace=True)

    with pytest.raises(IngestionServiceError):
        await async_test_ingest(service)

    assert service.statuses == ["processing", "failed"]


async def async_test_ingest(service: FakeIngestionService):
    return await service.ingest_document(
        document_id=DOCUMENT_ID,
        user_id=USER_ID,
        access_token="test-access-token",
    )
