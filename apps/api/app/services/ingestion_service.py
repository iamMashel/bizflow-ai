import logging
from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from typing import Any, cast
from uuid import UUID

import httpx
from docx import Document
from pypdf import PdfReader

from app.core.config import Settings, get_settings
from app.services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)

SUPPORTED_INGESTION_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150


class IngestionStage(StrEnum):
    DOWNLOAD_FILE = "download_file"
    EXTRACT_TEXT = "extract_text"
    CHUNK_TEXT = "chunk_text"
    EMBED_CHUNK = "embed_chunk"
    INSERT_CHUNKS = "insert_chunks"
    UPDATE_STATUS = "update_status"


class IngestionServiceError(Exception):
    def __init__(
        self,
        message: str,
        *,
        status_code: int = 502,
        stage: IngestionStage | None = None,
        supabase_status_code: int | None = None,
        supabase_body: str | None = None,
    ) -> None:
        super().__init__(message)
        self.status_code = status_code
        self.stage = stage
        self.supabase_status_code = supabase_status_code
        self.supabase_body = supabase_body


@dataclass(frozen=True)
class IngestedDocument:
    id: UUID
    filename: str
    storage_path: str


@dataclass(frozen=True)
class IngestionResult:
    document_id: UUID
    status: str
    chunks_created: int


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ExtractedText:
    text: str
    pages: list[ExtractedPage]


@dataclass(frozen=True)
class IngestionChunk:
    content: str
    metadata: dict[str, Any]


class IngestionService:
    def __init__(
        self,
        settings: Settings | None = None,
        embedding_service: EmbeddingService | None = None,
    ) -> None:
        self.settings = settings or get_settings()
        self.embedding_service = embedding_service or EmbeddingService(self.settings)
        self.supabase_url = self.settings.supabase_url.rstrip("/")
        self.supabase_anon_key = self.settings.supabase_anon_key
        self.service_role_key = self.settings.supabase_service_role_key
        self.bucket = self.settings.supabase_storage_bucket

    async def ingest_document(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
    ) -> IngestionResult:
        document = await self._get_document(
            document_id=document_id,
            user_id=user_id,
            access_token=access_token,
        )

        if self._extension(document.filename) not in SUPPORTED_INGESTION_EXTENSIONS:
            raise IngestionServiceError(
                "Only TXT, MD, DOCX, and PDF ingestion is currently supported.",
                status_code=400,
            )

        try:
            await self._update_document_status(
                document_id=document_id,
                user_id=user_id,
                access_token=access_token,
                status="processing",
            )
            content = await self._run_stage(
                document_id=document_id,
                stage=IngestionStage.DOWNLOAD_FILE,
                operation=lambda: self._download_storage_object(document.storage_path),
            )
            extracted_text = self._run_sync_stage(
                document_id=document_id,
                stage=IngestionStage.EXTRACT_TEXT,
                operation=lambda: self._extract_text(content, document.filename),
            )
            chunks = self._run_sync_stage(
                document_id=document_id,
                stage=IngestionStage.CHUNK_TEXT,
                operation=lambda: self.chunk_extracted_text(extracted_text),
            )
            await self._run_stage(
                document_id=document_id,
                stage=IngestionStage.INSERT_CHUNKS,
                operation=lambda: self._replace_chunks(
                    document_id=document_id,
                    user_id=user_id,
                    access_token=access_token,
                    chunks=chunks,
                ),
            )
            await self._update_document_status(
                document_id=document_id,
                user_id=user_id,
                access_token=access_token,
                status="completed",
            )
            return IngestionResult(
                document_id=document_id,
                status="completed",
                chunks_created=len(chunks),
            )
        except Exception as exc:
            self._log_ingestion_exception(
                document_id=document_id,
                stage=self._stage_for_exception(exc),
                exc=exc,
            )
            await self._mark_failed(
                document_id=document_id,
                user_id=user_id,
                access_token=access_token,
                reason=self._safe_failure_reason(exc),
            )
            raise

    async def _run_stage(
        self,
        *,
        document_id: UUID,
        stage: IngestionStage,
        operation: Any,
    ) -> Any:
        try:
            return await operation()
        except Exception as exc:
            self._log_ingestion_exception(document_id=document_id, stage=stage, exc=exc)
            raise self._ensure_stage(exc, stage) from exc

    def _run_sync_stage(
        self,
        *,
        document_id: UUID,
        stage: IngestionStage,
        operation: Any,
    ) -> Any:
        try:
            return operation()
        except Exception as exc:
            self._log_ingestion_exception(document_id=document_id, stage=stage, exc=exc)
            raise self._ensure_stage(exc, stage) from exc

    async def _get_document(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
    ) -> IngestedDocument:
        self._require_supabase_config()
        async with httpx.AsyncClient(timeout=20.0) as client:
            response = await client.get(
                f"{self.supabase_url}/rest/v1/documents",
                headers=self._database_headers(access_token),
                params={
                    "select": "id,filename,storage_path",
                    "id": f"eq.{document_id}",
                    "user_id": f"eq.{user_id}",
                    "limit": "1",
                },
            )

        self._raise_for_supabase_error(response, "Unable to read document.")
        payload = response.json()
        if not isinstance(payload, list) or not payload:
            raise IngestionServiceError("Document not found.", status_code=404)

        row = cast(dict[str, Any], payload[0])
        filename = row.get("filename")
        storage_path = row.get("storage_path")
        if not isinstance(filename, str) or not isinstance(storage_path, str):
            raise IngestionServiceError("Document is missing storage metadata.")

        return IngestedDocument(id=document_id, filename=filename, storage_path=storage_path)

    async def _download_storage_object(self, storage_path: str) -> bytes:
        self._require_supabase_config()
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.get(
                f"{self.supabase_url}/storage/v1/object/{self.bucket}/{storage_path}",
                headers=self._service_headers(),
            )

        self._raise_for_supabase_error(
            response,
            "Unable to download document.",
            stage=IngestionStage.DOWNLOAD_FILE,
        )
        return response.content

    async def _replace_chunks(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
        chunks: list[IngestionChunk],
    ) -> None:
        await self._delete_existing_chunks(
            document_id=document_id,
            user_id=user_id,
            access_token=access_token,
        )
        if not chunks:
            return

        rows = []
        for index, chunk in enumerate(chunks):
            rows.append(
                {
                    "document_id": str(document_id),
                    "user_id": str(user_id),
                    "chunk_index": index,
                    "content": chunk.content,
                    "token_count": self._estimate_token_count(chunk.content),
                    "metadata": chunk.metadata,
                    "embedding": self._embed_chunk(document_id=document_id, chunk=chunk.content),
                }
            )

        async def insert_chunks() -> None:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{self.supabase_url}/rest/v1/document_chunks",
                    headers={
                        **self._database_headers(access_token),
                        "Content-Type": "application/json",
                        "Prefer": "return=minimal",
                    },
                    json=rows,
                )

            self._raise_for_supabase_error(
                response,
                "Unable to store document chunks.",
                stage=IngestionStage.INSERT_CHUNKS,
            )

        await self._run_stage(
            document_id=document_id,
            stage=IngestionStage.INSERT_CHUNKS,
            operation=insert_chunks,
        )

    async def _delete_existing_chunks(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
    ) -> None:
        async with httpx.AsyncClient(timeout=20.0) as client:
            response = await client.delete(
                f"{self.supabase_url}/rest/v1/document_chunks",
                headers=self._database_headers(access_token),
                params={
                    "document_id": f"eq.{document_id}",
                    "user_id": f"eq.{user_id}",
                },
            )

        self._raise_for_supabase_error(
            response,
            "Unable to clear existing document chunks.",
            stage=IngestionStage.INSERT_CHUNKS,
        )

    async def _update_document_status(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
        status: str,
        metadata_patch: dict[str, Any] | None = None,
    ) -> None:
        payload: dict[str, Any] = {"status": status}
        if metadata_patch is not None:
            payload["metadata"] = metadata_patch

        async with httpx.AsyncClient(timeout=20.0) as client:
            response = await client.patch(
                f"{self.supabase_url}/rest/v1/documents",
                headers={
                    **self._database_headers(access_token),
                    "Content-Type": "application/json",
                },
                params={
                    "id": f"eq.{document_id}",
                    "user_id": f"eq.{user_id}",
                },
                json=payload,
            )

        self._raise_for_supabase_error(
            response,
            "Unable to update document status.",
            stage=IngestionStage.UPDATE_STATUS,
        )

    async def _mark_failed(
        self,
        *,
        document_id: UUID,
        user_id: UUID,
        access_token: str,
        reason: str,
    ) -> None:
        try:
            await self._update_document_status(
                document_id=document_id,
                user_id=user_id,
                access_token=access_token,
                status="failed",
                metadata_patch={"ingestion_error": reason},
            )
        except IngestionServiceError:
            logger.warning(
                "Unable to mark document ingestion failed: document_id=%s",
                document_id,
                exc_info=True,
            )

    def _database_headers(self, access_token: str) -> dict[str, str]:
        if not self.supabase_anon_key or not access_token:
            raise IngestionServiceError("Supabase database access is not configured.")

        return {
            "apikey": self.supabase_anon_key,
            "Authorization": f"Bearer {access_token}",
        }

    def _service_headers(self) -> dict[str, str]:
        return {
            "apikey": self.service_role_key,
            "Authorization": f"Bearer {self.service_role_key}",
        }

    def _require_supabase_config(self) -> None:
        if not self.supabase_url or not self.service_role_key:
            raise IngestionServiceError("Supabase ingestion storage is not configured.")

    @staticmethod
    def _extract_text(content: bytes, filename: str) -> ExtractedText:
        extension = IngestionService._extension(filename)
        if extension == ".docx":
            return ExtractedText(
                text=IngestionService._extract_docx_text(content),
                pages=[],
            )
        if extension == ".pdf":
            return IngestionService._extract_pdf_text(content)

        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise IngestionServiceError(
                "Uploaded text document must be valid UTF-8.",
                status_code=400,
            ) from exc
        return ExtractedText(text=text, pages=[])

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise IngestionServiceError(
                "Uploaded DOCX document could not be read.",
                status_code=400,
            ) from exc

        text_parts = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(part for part in text_parts if part)

    @staticmethod
    def _extract_pdf_text(content: bytes) -> ExtractedText:
        try:
            reader = PdfReader(BytesIO(content))
        except Exception as exc:
            raise IngestionServiceError(
                "Uploaded PDF document could not be read.",
                status_code=400,
            ) from exc

        pages: list[ExtractedPage] = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(ExtractedPage(page_number=index, text=page_text))

        return ExtractedText(
            text="\n\n".join(page.text for page in pages),
            pages=pages,
        )

    @staticmethod
    def chunk_extracted_text(extracted_text: ExtractedText) -> list[IngestionChunk]:
        if extracted_text.pages:
            chunks: list[IngestionChunk] = []
            for page in extracted_text.pages:
                chunks.extend(
                    IngestionChunk(
                        content=chunk,
                        metadata={"page_number": page.page_number},
                    )
                    for chunk in IngestionService.chunk_text(page.text)
                )
            return chunks

        return [
            IngestionChunk(content=chunk, metadata={})
            for chunk in IngestionService.chunk_text(extracted_text.text)
        ]

    @staticmethod
    def chunk_text(
        text: str,
        *,
        chunk_size: int = CHUNK_SIZE,
        overlap: int = CHUNK_OVERLAP,
    ) -> list[str]:
        normalized = text.strip()
        if not normalized:
            return []

        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(start + chunk_size, len(normalized))
            chunks.append(normalized[start:end])
            if end == len(normalized):
                break
            start = max(end - overlap, start + 1)

        return chunks

    @staticmethod
    def _estimate_token_count(text: str) -> int:
        return max(1, len(text.split()))

    @staticmethod
    def _extension(filename: str) -> str:
        if "." not in filename:
            return ""
        return f".{filename.rsplit('.', 1)[-1].lower()}"

    @staticmethod
    def _raise_for_supabase_error(
        response: httpx.Response,
        message: str,
        *,
        stage: IngestionStage | None = None,
    ) -> None:
        if response.status_code >= 400:
            logger.warning(
                "Supabase REST error: status_code=%s path=%s body=%s stage=%s",
                response.status_code,
                response.request.url.path,
                response.text,
                stage,
            )
            raise IngestionServiceError(
                message,
                stage=stage,
                supabase_status_code=response.status_code,
                supabase_body=response.text,
            )

    def _embed_chunk(self, *, document_id: UUID, chunk: str) -> list[float]:
        result = self._run_sync_stage(
            document_id=document_id,
            stage=IngestionStage.EMBED_CHUNK,
            operation=lambda: self.embedding_service.embed_text(chunk),
        )
        return cast(list[float], result)

    @staticmethod
    def _ensure_stage(exc: Exception, stage: IngestionStage) -> IngestionServiceError:
        if isinstance(exc, IngestionServiceError):
            if exc.stage is None:
                exc.stage = stage
            return exc
        return IngestionServiceError(str(exc), stage=stage)

    @staticmethod
    def _stage_for_exception(exc: Exception) -> IngestionStage:
        if isinstance(exc, IngestionServiceError) and exc.stage is not None:
            return exc.stage
        return IngestionStage.UPDATE_STATUS

    @staticmethod
    def _safe_failure_reason(exc: Exception) -> str:
        message = str(exc) or exc.__class__.__name__
        return message[:500]

    @staticmethod
    def _log_ingestion_exception(
        *,
        document_id: UUID,
        stage: IngestionStage,
        exc: Exception,
    ) -> None:
        supabase_status_code = None
        supabase_body = None
        if isinstance(exc, IngestionServiceError):
            supabase_status_code = exc.supabase_status_code
            supabase_body = exc.supabase_body

        logger.warning(
            (
                "Document ingestion error: document_id=%s stage=%s "
                "exception_type=%s exception_message=%s supabase_status=%s supabase_body=%s"
            ),
            document_id,
            stage,
            exc.__class__.__name__,
            str(exc),
            supabase_status_code,
            supabase_body,
        )
